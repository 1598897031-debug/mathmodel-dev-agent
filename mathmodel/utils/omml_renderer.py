"""
OMML (Office Math Markup Language) Renderer
Converts LaTeX formulas to native Word equation objects.

Usage:
    from mathmodel.utils.omml_renderer import insert_formula, insert_display_formula
    insert_formula(paragraph, "d_i = \\sqrt{(x_{s_i}-x_n)^2 + y_n^2}")
"""

import re
from lxml import etree

M = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def _xml(tag, children="", attrs=None):
    """Build an OMML XML element string."""
    attr_str = ""
    if attrs:
        attr_str = " " + " ".join(f'm:{k}="{v}"' for k, v in attrs.items())
    if children:
        return f"<m:{tag}{attr_str}>{children}</m:{tag}>"
    return f"<m:{tag}{attr_str}/>"


def _r(text, font="Cambria Math", sz="24", italic=None):
    """Build an OMML run element."""
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    sty = ""
    if italic is not None:
        sty_val = "i" if italic else "p"
        sty = f"<m:sty m:val=\"{sty_val}\"/>"
    return (
        f"<m:r>"
        f"<m:rPr>{sty}"
        f"<m:rFonts m:ascii=\"{font}\" m:hAnsi=\"{font}\" m:eastAsia=\"{font}\"/>"
        f"<m:sz m:val=\"{sz}\"/><m:szCs m:val=\"{sz}\"/>"
        f"</m:rPr>"
        f"<m:t xml:space=\"preserve\">{text}</m:t>"
        f"</m:r>"
    )


def _frac(num_xml, den_xml):
    """Build a fraction."""
    return f"<m:f><m:fNum>{num_xml}</m:fNum><m:fDen>{den_xml}</m:fDen></m:f>"


def _rad(content_xml, deg_xml=None):
    """Build a radical (square root)."""
    if deg_xml:
        return f"<m:rad><m:radPr><m:degHide m:val=\"0\"/></m:radPr><m:deg>{deg_xml}</m:deg><m:e>{content_xml}</m:e></m:rad>"
    return f"<m:rad><m:radPr><m:degHide m:val=\"1\"/></m:radPr><m:e>{content_xml}</m:e></m:rad>"


def _sub(base_xml, sub_xml):
    """Build subscript."""
    return f"<m:sSub><m:e>{base_xml}</m:e><m:sub>{sub_xml}</m:sub></m:sSub>"


def _sup(base_xml, sup_xml):
    """Build superscript."""
    return f"<m:sSup><m:e>{base_xml}</m:e><m:sup>{sup_xml}</m:sup></m:sSup>"


def _subsup(base_xml, sub_xml, sup_xml):
    """Build sub-superscript."""
    return f"<m:sSubSup><m:e>{base_xml}</m:e><m:sub>{sub_xml}</m:sub><m:sup>{sup_xml}</m:sup></m:sSubSup>"


def _nary(op, sub_xml, sup_xml, body_xml):
    """Build n-ary operator (sum, prod, int)."""
    parts = f"<m:naryPr><m:chr m:val=\"{op}\"/><m:limLoc m:val=\"undOvr\"/></m:naryPr>"
    if sub_xml:
        parts += f"<m:sub>{sub_xml}</m:m:sub>" if not sub_xml.startswith("<m:") else f"<m:sub>{sub_xml}</m:sub>"
    if sup_xml:
        parts += f"<m:sup>{sup_xml}</m:sup>"
    parts += f"<m:e>{body_xml}</m:e>"
    return f"<m:nary>{parts}</m:nary>"


# Greek letter mapping
_GREEK = {
    "alpha": "α", "beta": "β", "gamma": "γ", "delta": "δ",
    "epsilon": "ε", "zeta": "ζ", "eta": "η", "theta": "θ",
    "lambda": "λ", "mu": "μ", "nu": "ν", "xi": "ξ",
    "pi": "π", "rho": "ρ", "sigma": "σ", "tau": "τ",
    "phi": "φ", "chi": "χ", "psi": "ψ", "omega": "ω",
    "Gamma": "Γ", "Delta": "Δ", "Theta": "Θ", "Lambda": "Λ",
    "Sigma": "Σ", "Phi": "Φ", "Psi": "Ψ", "Omega": "Ω",
}

_OPS = {
    "cdot": "·", "times": "×", "div": "÷",
    "pm": "±", "mp": "∓",
    "leq": "≤", "geq": "≥", "neq": "≠",
    "approx": "≈", "equiv": "≡",
    "leftarrow": "←", "rightarrow": "→",
    "Rightarrow": "⇒", "Leftarrow": "⇐",
    "infty": "∞", "partial": "∂", "nabla": "∇",
    "quad": " ", "qquad": "  ",
}


def _parse_latex(tex, sz="24"):
    """Parse LaTeX string into OMML XML string."""
    tex = tex.strip()
    if tex.startswith("$$") and tex.endswith("$$"):
        tex = tex[2:-2].strip()
    elif tex.startswith("$") and tex.endswith("$"):
        tex = tex[1:-1].strip()

    # Remove style commands
    tex = re.sub(r'\\(?:display|text)style\s*', '', tex)

    parts = []
    i = 0
    while i < len(tex):
        c = tex[i]

        # LaTeX command
        if c == '\\':
            m = re.match(r'\\([a-zA-Z]+)\s*', tex[i:])
            if m:
                cmd = m.group(1)
                i += m.end()

                if cmd == "frac":
                    num, i = _extract_brace(tex, i)
                    den, i = _extract_brace(tex, i)
                    parts.append(_frac(_parse_latex(num, sz), _parse_latex(den, sz)))

                elif cmd == "sqrt":
                    deg = None
                    if i < len(tex) and tex[i] == '[':
                        deg, i = _extract_bracket(tex, i)
                    content, i = _extract_brace(tex, i)
                    deg_xml = _parse_latex(deg, sz) if deg else None
                    parts.append(_rad(_parse_latex(content, sz), deg_xml))

                elif cmd in ("sum", "prod", "int"):
                    op_char = {"sum": "∑", "prod": "∏", "int": "∫"}[cmd]
                    sub_xml = sup_xml = ""
                    if i < len(tex) and tex[i] == '_':
                        i += 1
                        sub, i = _extract_single_or_brace(tex, i)
                        sub_xml = _parse_latex(sub, sz)
                    if i < len(tex) and tex[i] == '^':
                        i += 1
                        sup, i = _extract_single_or_brace(tex, i)
                        sup_xml = _parse_latex(sup, sz)
                    parts.append(_nary(op_char, sub_xml, sup_xml, ""))

                elif cmd == "text":
                    content, i = _extract_brace(tex, i)
                    parts.append(_r(content, font="宋体", sz=sz))

                elif cmd == "mathbf":
                    content, i = _extract_brace(tex, i)
                    parts.append(_parse_latex(content, sz))

                elif cmd in _GREEK:
                    parts.append(_r(_GREEK[cmd], sz=sz, italic=True))

                elif cmd in _OPS:
                    parts.append(_r(_OPS[cmd], sz=sz))

                elif cmd in ("left", "right"):
                    if i < len(tex):
                        delim = tex[i]
                        if delim in "()[]|.":
                            parts.append(_r(delim, sz=sz))
                        i += 1

                elif cmd in ("quad", "qquad"):
                    parts.append(_r(" ", sz=sz))

                elif cmd in ("sin", "cos", "tan", "log", "ln", "exp", "lim", "max", "min", "argmin", "argmax"):
                    parts.append(_r(cmd, sz=sz, italic=False))

                else:
                    parts.append(_r(f"\\{cmd}", sz=sz))
            else:
                i += 1
                if i < len(tex):
                    i += 1
            continue

        # Subscript
        if c == '_':
            i += 1
            sub, i = _extract_single_or_brace(tex, i)
            sup = None
            if i < len(tex) and tex[i] == '^':
                i += 1
                sup, i = _extract_single_or_brace(tex, i)
            # Get last element as base
            if parts:
                base = parts.pop()
            else:
                base = _r("", sz=sz)
            if sup:
                parts.append(_subsup(base, _parse_latex(sub, sz), _parse_latex(sup, sz)))
            else:
                parts.append(_sub(base, _parse_latex(sub, sz)))
            continue

        # Superscript
        if c == '^':
            i += 1
            sup, i = _extract_single_or_brace(tex, i)
            if parts:
                base = parts.pop()
            else:
                base = _r("", sz=sz)
            parts.append(_sup(base, _parse_latex(sup, sz)))
            continue

        # Brace group
        if c == '{':
            content, i = _extract_brace(tex, i)
            parts.append(_parse_latex(content, sz))
            continue

        # Whitespace
        if c in " \t":
            i += 1
            continue

        # Numbers
        if c.isdigit():
            nm = re.match(r'\d+\.?\d*', tex[i:])
            if nm:
                parts.append(_r(nm.group(), sz=sz))
                i += nm.end()
                continue

        # Regular character
        parts.append(_r(c, sz=sz, italic=True if c.isalpha() else None))
        i += 1

    return "".join(parts)


def _extract_brace(tex, i):
    if i >= len(tex) or tex[i] != '{':
        return "", i
    depth = 1
    start = i + 1
    i += 1
    while i < len(tex) and depth > 0:
        if tex[i] == '{': depth += 1
        elif tex[i] == '}': depth -= 1
        i += 1
    return tex[start:i-1], i


def _extract_single_or_brace(tex, i):
    if i >= len(tex):
        return "", i
    if tex[i] == '{':
        return _extract_brace(tex, i)
    return tex[i], i + 1


def _extract_bracket(tex, i):
    if i >= len(tex) or tex[i] != '[':
        return "", i
    start = i + 1
    i += 1
    while i < len(tex) and tex[i] != ']':
        i += 1
    return tex[start:i], i + 1


def insert_formula(paragraph, latex, font_size_pt=12, display=False):
    """
    Insert a LaTeX formula as a native Word OMML equation object.

    Args:
        paragraph: python-docx Paragraph
        latex: LaTeX string
        font_size_pt: font size in points
        display: if True, render as display equation (centered, own line)
    """
    sz = str(font_size_pt * 2)  # half-points

    try:
        inner_xml = _parse_latex(latex, sz=sz)
    except Exception:
        # Fallback to plain text
        run = paragraph.add_run(latex)
        run.font.name = 'Cambria Math'
        run.font.size = paragraph.part.element.nsmap and __import__('docx.shared', fromlist=['Pt']).Pt(font_size_pt)
        return

    if display:
        omml_xml = f'<m:oMathPara xmlns:m="{M}"><m:oMathParaPr><m:jc m:val="center"/></m:oMathParaPr><m:oMath>{inner_xml}</m:oMath></m:oMathPara>'
    else:
        omml_xml = f'<m:oMath xmlns:m="{M}">{inner_xml}</m:oMath>'

    try:
        elem = etree.fromstring(omml_xml.encode('utf-8'))
        run = paragraph.add_run()
        run._element.append(elem)
    except Exception:
        # Fallback to plain text
        run = paragraph.add_run(latex)
        run.font.name = 'Cambria Math'
        from docx.shared import Pt
        run.font.size = Pt(font_size_pt)


def insert_display_formula(paragraph, latex, font_size_pt=12):
    """Insert a display (centered, standalone) formula."""
    insert_formula(paragraph, latex, font_size_pt, display=True)


def insert_inline_formula(paragraph, latex, font_size_pt=12):
    """Insert an inline formula."""
    insert_formula(paragraph, latex, font_size_pt, display=False)
